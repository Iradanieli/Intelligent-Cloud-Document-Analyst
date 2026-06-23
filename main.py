from datetime import datetime, timezone
import io
import fitz  # PyMuPDF
import docx  # python-docx
from typing import Any, Dict, List, Literal, Tuple
from uuid import uuid4
from fastapi import FastAPI, UploadFile, File, HTTPException
from pydantic import BaseModel, Field, field_validator


app = FastAPI(
    title="Business Document Metadata API",
    description="Metadata enrichment service for the n8n Business Document Analyst workflow.",
    version="1.0.0",
)


Category = Literal[
    "invoice",
    "contract",
    "purchase_order",
    "vendor_agreement",
    "report",
    "other",
]

Sentiment = Literal["positive", "neutral", "negative"]
SensitivityLevel = Literal["public", "internal", "confidential"]
RoutingTag = Literal["auto-approved", "needs-review", "escalate"]

CATEGORIES: List[Category] = [
    "invoice",
    "contract",
    "purchase_order",
    "vendor_agreement",
    "report",
    "other",
]

DEPARTMENT_BY_CATEGORY: Dict[str, str] = {
    "invoice": "Finance",
    "contract": "Legal",
    "purchase_order": "Procurement",
    "vendor_agreement": "Vendor Management",
    "report": "Management",
    "other": "General",
}

CONFIDENTIAL_KEYWORDS = {
    "confidential",
    "nda",
    "non-disclosure",
    "non disclosure",
    "bank account",
    "tax id",
    "ssn",
    "social security",
    "salary",
    "wire transfer",
}

INTERNAL_KEYWORDS = {
    "invoice",
    "contract",
    "purchase order",
    "purchase_order",
    "vendor agreement",
    "vendor_agreement",
    "payment terms",
    "budget",
    "quote",
}

KEY_BUSINESS_ENTITY_GROUPS = {
    "organizations",
    "dates",
    "amounts",
    "invoice_numbers",
    "purchase_order_numbers",
    "contract_terms",
}


class Entities(BaseModel):
    people: List[str] = Field(default_factory=list)
    organizations: List[str] = Field(default_factory=list)
    dates: List[str] = Field(default_factory=list)
    amounts: List[str] = Field(default_factory=list)
    invoice_numbers: List[str] = Field(default_factory=list)
    purchase_order_numbers: List[str] = Field(default_factory=list)
    contract_terms: List[str] = Field(default_factory=list)

    model_config = {"extra": "allow"}


class SensitivityRequest(BaseModel):
    text: str = ""
    entities: Entities = Field(default_factory=Entities)


class SensitivityResponse(BaseModel):
    sensitivity: SensitivityLevel
    matched_keywords: List[str]


class GeminiResult(BaseModel):
    classification: Category
    sentiment: Sentiment = "neutral"
    confidence_score: float = Field(ge=0.0, le=1.0)
    summary: str = ""
    entities: Entities = Field(default_factory=Entities)
    action_items: List[str] = Field(default_factory=list)
    document_text: str = ""

    @field_validator("classification", mode="before")
    @classmethod
    def normalize_classification(cls, value: Any) -> str:
        normalized = str(value or "other").strip().lower().replace(" ", "_").replace("-", "_")
        return normalized if normalized in CATEGORIES else "other"

    @field_validator("sentiment", mode="before")
    @classmethod
    def normalize_sentiment(cls, value: Any) -> str:
        normalized = str(value or "neutral").strip().lower()
        return normalized if normalized in {"positive", "neutral", "negative"} else "neutral"


class EnrichResponse(BaseModel):
    document_id: str
    processed_at: str
    department: str
    sensitivity: SensitivityLevel
    routing_tag: RoutingTag
    business_tags: List[str]
    adjusted_confidence_score: float


def flatten_values(value: Any) -> List[str]:
    if value is None:
        return []
    if isinstance(value, str):
        return [value]
    if isinstance(value, dict):
        values: List[str] = []
        for item in value.values():
            values.extend(flatten_values(item))
        return values
    if isinstance(value, (list, tuple, set)):
        values = []
        for item in value:
            values.extend(flatten_values(item))
        return values
    return [str(value)]


def searchable_text(text: str, entities: Entities) -> str:
    entity_text = " ".join(flatten_values(entities.model_dump()))
    return f"{text} {entity_text}".lower()


def classify_sensitivity(text: str, entities: Entities) -> Tuple[SensitivityLevel, List[str]]:
    content = searchable_text(text, entities)
    matched_confidential = sorted(keyword for keyword in CONFIDENTIAL_KEYWORDS if keyword in content)
    if matched_confidential:
        return "confidential", matched_confidential

    matched_internal = sorted(keyword for keyword in INTERNAL_KEYWORDS if keyword in content)
    if matched_internal:
        return "internal", matched_internal

    return "public", []


def has_key_entities(entities: Entities) -> bool:
    entity_data = entities.model_dump()
    present_groups = [
        group
        for group in KEY_BUSINESS_ENTITY_GROUPS
        if flatten_values(entity_data.get(group))
    ]
    return len(present_groups) >= 2


def has_incomplete_entities(entities: Entities) -> bool:
    entity_data = entities.model_dump()
    return not flatten_values(entity_data)


def adjusted_confidence(data: GeminiResult, sensitivity: SensitivityLevel) -> float:
    score = data.confidence_score
    if has_incomplete_entities(data.entities):
        score -= 0.15
    if sensitivity == "confidential":
        score -= 0.10
    if has_key_entities(data.entities):
        score += 0.05
    return round(min(max(score, 0.0), 1.0), 2)


def routing_tag(data: GeminiResult, sensitivity: SensitivityLevel, score: float) -> RoutingTag:
    if sensitivity == "confidential" or data.sentiment == "negative":
        return "escalate"
    if score < 0.70 or has_incomplete_entities(data.entities):
        return "needs-review"
    return "auto-approved"


def business_tags(data: GeminiResult, sensitivity: SensitivityLevel, route: RoutingTag) -> List[str]:
    tags = {data.classification, sensitivity, route}
    if data.sentiment != "neutral":
        tags.add(f"sentiment-{data.sentiment}")
    if data.action_items:
        tags.add("has-action-items")
    if data.entities.amounts:
        tags.add("has-amounts")
    return sorted(tags)


@app.get("/health")
def health() -> Dict[str, str]:
    return {"status": "ok", "service": "business-document-metadata-api"}


@app.get("/categories")
def categories() -> Dict[str, List[str]]:
    return {"categories": CATEGORIES}


@app.post("/sensitivity", response_model=SensitivityResponse)
def sensitivity(request: SensitivityRequest) -> SensitivityResponse:
    level, matched_keywords = classify_sensitivity(request.text, request.entities)
    return SensitivityResponse(sensitivity=level, matched_keywords=matched_keywords)


@app.post("/enrich", response_model=EnrichResponse)
def enrich(data: GeminiResult) -> EnrichResponse:
    sensitivity_level, _ = classify_sensitivity(
        " ".join([data.summary, data.document_text]),
        data.entities,
    )
    score = adjusted_confidence(data, sensitivity_level)
    route = routing_tag(data, sensitivity_level, score)

    return EnrichResponse(
        document_id=str(uuid4()),
        processed_at=datetime.now(timezone.utc).isoformat(),
        department=DEPARTMENT_BY_CATEGORY[data.classification],
        sensitivity=sensitivity_level,
        routing_tag=route,
        business_tags=business_tags(data, sensitivity_level, route),
        adjusted_confidence_score=score,
    )

@app.post("/extract")
async def extract_text(file: UploadFile = File(...)):
    # Safely get the file extension
    filename = file.filename or "unknown.txt"
    ext = filename.split(".")[-1].lower()
    
    # Read the file bytes into memory
    content = await file.read()
    extracted_text = ""

    try:
        if ext == "pdf":
            # Extract PDF text using PyMuPDF
            doc = fitz.Document(stream=content, filetype="pdf")
            for page in doc:
                extracted_text += page.get_text() + "\n"
            doc.close()
            
        elif ext == "docx":
            # Extract DOCX text using python-docx
            doc = docx.Document(io.BytesIO(content))
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
                
        elif ext in ["txt", "csv", "md"]:
            # Standard text decoding
            extracted_text = content.decode("utf-8")
            
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error parsing file: {str(e)}")

    return {
        "filename": filename,
        "file_type": ext,
        "text": extracted_text.strip()
    }
